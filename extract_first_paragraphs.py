from docx import Document

doc = Document("asl\u0131 s\u0131la \u00f6zate\u015f tez.docx")

for i in range(120):
    if i < len(doc.paragraphs):
        p = doc.paragraphs[i]
        print(f"P{i} ({p.style.name}): {p.text}")
